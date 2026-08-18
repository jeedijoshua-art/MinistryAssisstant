from uuid import UUID, uuid4
from sqlalchemy.orm import Session
from sqlalchemy import select, desc
from typing import List, Optional, Dict, Any
from app.models.domain import Sermon, SermonSection, SermonHistory, SermonTemplate, Pastor

class SermonService:
    def __init__(self, db: Session):
        self.db = db

    def get_sermons(self, pastor_id: Optional[UUID] = None, limit: int = 20, offset: int = 0) -> List[Sermon]:
        query = select(Sermon).order_by(desc(Sermon.created_at))
        if pastor_id:
            query = query.where(Sermon.pastor_id == pastor_id)
        
        result = self.db.execute(query.limit(limit).offset(offset))
        return list(result.scalars().all())

    def get_sermon(self, sermon_id: UUID) -> Optional[Sermon]:
        return self.db.get(Sermon, sermon_id)

    def create_sermon(self, data: Dict[str, Any], pastor_id: Optional[UUID] = None) -> Sermon:
        sermon = Sermon(
            id=uuid4(),
            pastor_id=pastor_id,
            title=data.get("title", "Untitled Sermon"),
            theme=data.get("theme"),
            main_verse=data.get("main_verse"),
            bible_version=data.get("bible_version", "KJV"),
            audience=data.get("audience"),
            occasion=data.get("occasion"),
            estimated_duration=data.get("estimated_duration"),
            date_preached=data.get("date_preached"),
            content=data.get("content"),
            status=data.get("status", "draft"),
            notes=data.get("notes")
        )
        self.db.add(sermon)
        
        if "sections" in data:
            for i, sec in enumerate(data["sections"]):
                section = SermonSection(
                    id=uuid4(),
                    sermon_id=sermon.id,
                    title=sec.get("title", ""),
                    content=sec.get("content", ""),
                    order_index=sec.get("order_index", i)
                )
                self.db.add(section)
                
        self.db.commit()
        self.db.refresh(sermon)
        return sermon

    def update_sermon(self, sermon_id: UUID, data: Dict[str, Any]) -> Optional[Sermon]:
        sermon = self.get_sermon(sermon_id)
        if not sermon:
            return None
            
        # Snapshot for history if content is changing significantly
        if "content" in data and sermon.content != data["content"]:
            self._save_history(sermon.id, sermon.content, "Auto-saved before update")

        for key, value in data.items():
            if hasattr(sermon, key) and key not in ["id", "sections", "history", "tags", "references"]:
                setattr(sermon, key, value)
                
        # Handle sections update
        if "sections" in data:
            # Delete old sections and recreate for simplicity, or update by ID
            # For simplicity in this sprint, we'll recreate them
            for old_sec in sermon.sections:
                self.db.delete(old_sec)
            for i, sec in enumerate(data["sections"]):
                new_sec = SermonSection(
                    id=uuid4(),
                    sermon_id=sermon.id,
                    title=sec.get("title", ""),
                    content=sec.get("content", ""),
                    order_index=sec.get("order_index", i)
                )
                self.db.add(new_sec)

        self.db.commit()
        self.db.refresh(sermon)
        return sermon

    def delete_sermon(self, sermon_id: UUID) -> bool:
        sermon = self.get_sermon(sermon_id)
        if sermon:
            self.db.delete(sermon)
            self.db.commit()
            return True
        return False

    def _save_history(self, sermon_id: UUID, content: str, note: str = ""):
        if not content:
            return
        history = SermonHistory(
            id=uuid4(),
            sermon_id=sermon_id,
            content_snapshot=content,
            version_note=note
        )
        self.db.add(history)

    def get_history(self, sermon_id: UUID) -> List[SermonHistory]:
        query = select(SermonHistory).where(SermonHistory.sermon_id == sermon_id).order_by(desc(SermonHistory.created_at))
        result = self.db.execute(query)
        return list(result.scalars().all())

    def get_templates(self) -> List[SermonTemplate]:
        query = select(SermonTemplate).order_by(SermonTemplate.name)
        result = self.db.execute(query)
        return list(result.scalars().all())

    # Export functionality stub
    def export_sermon_docx(self, sermon_id: UUID) -> str:
        sermon = self.get_sermon(sermon_id)
        if not sermon:
            raise ValueError("Sermon not found")
            
        import docx
        doc = docx.Document()
        doc.add_heading(sermon.title, 0)
        if sermon.main_verse:
            doc.add_paragraph(f"Main Text: {sermon.main_verse} ({sermon.bible_version})")
            
        doc.add_paragraph(sermon.content or "")
        
        for section in sorted(sermon.sections, key=lambda x: x.order_index):
            doc.add_heading(section.title, level=2)
            doc.add_paragraph(section.content)
            
        filepath = f"/tmp/{sermon.id}.docx"
        doc.save(filepath)
        return filepath
